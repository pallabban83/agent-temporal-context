"""Multimodal document parser for PDFs and DOCX with images."""

import io
import logging
import hashlib
from typing import List, Dict, Any, Tuple
from PIL import Image
from pypdf import PdfReader
from docx import Document
from google.cloud import vision
from google.cloud import storage

logger = logging.getLogger(__name__)


class MultimodalDocumentParser:
    """Parse documents with both text and images."""

    def __init__(self, project_id: str, gcs_bucket: str):
        """Initialize the multimodal parser.

        Args:
            project_id: Google Cloud project ID
            gcs_bucket: GCS bucket for storing images
        """
        self.project_id = project_id
        self.gcs_bucket = gcs_bucket
        self.vision_client = vision.ImageAnnotatorClient()
        self.storage_client = storage.Client(project=project_id)

    def extract_images_from_pdf(self, pdf_bytes: bytes) -> List[Tuple[int, Image.Image]]:
        """Extract images from PDF.

        Args:
            pdf_bytes: PDF file content as bytes

        Returns:
            List of (page_number, image) tuples
        """
        pdf_file = io.BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)

        images = []
        for page_num, page in enumerate(reader.pages):
            if '/XObject' in page['/Resources']:
                xObject = page['/Resources']['/XObject'].get_object()

                for obj in xObject:
                    if xObject[obj]['/Subtype'] == '/Image':
                        try:
                            # Extract image data
                            size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                            data = xObject[obj].get_data()

                            # Determine color mode
                            if xObject[obj]['/ColorSpace'] == '/DeviceRGB':
                                mode = "RGB"
                            else:
                                mode = "P"

                            img = Image.frombytes(mode, size, data)
                            images.append((page_num, img))
                            logger.info(f"Extracted image from page {page_num + 1}: {size[0]}x{size[1]}")
                        except Exception as e:
                            logger.warning(f"Could not extract image from page {page_num}: {e}")
                            continue

        logger.info(f"Extracted {len(images)} images from PDF")
        return images

    def extract_images_from_docx(self, docx_bytes: bytes) -> List[Tuple[int, Image.Image]]:
        """Extract images from DOCX.

        Args:
            docx_bytes: DOCX file content as bytes

        Returns:
            List of (paragraph_index, image) tuples
        """
        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)

        images = []
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data))
                    images.append((i, img))
                    logger.info(f"Extracted image {i}: {img.size[0]}x{img.size[1]}")
                except Exception as e:
                    logger.warning(f"Could not extract image {i}: {e}")
                    continue

        logger.info(f"Extracted {len(images)} images from DOCX")
        return images

    def describe_image(self, image: Image.Image) -> Dict[str, Any]:
        """Generate description of image using Vision API.

        Args:
            image: PIL Image object

        Returns:
            Dictionary with image description and metadata
        """
        try:
            # Convert PIL Image to bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_byte_arr = img_byte_arr.getvalue()

            vision_image = vision.Image(content=img_byte_arr)

            # Detect labels (what's in the image)
            label_response = self.vision_client.label_detection(image=vision_image)
            labels = [label.description for label in label_response.label_annotations[:10]]

            # Detect text in image (OCR)
            text_response = self.vision_client.text_detection(image=vision_image)
            ocr_text = ""
            if text_response.text_annotations:
                ocr_text = text_response.text_annotations[0].description

            # Detect objects
            object_response = self.vision_client.object_localization(image=vision_image)
            objects = [obj.name for obj in object_response.localized_object_annotations[:5]]

            # Build comprehensive description
            description_parts = []

            if labels:
                description_parts.append(f"Image contains: {', '.join(labels)}")

            if objects:
                description_parts.append(f"Detected objects: {', '.join(objects)}")

            if ocr_text:
                # Limit OCR text to 500 characters
                ocr_preview = ocr_text[:500] + "..." if len(ocr_text) > 500 else ocr_text
                description_parts.append(f"Text in image: {ocr_preview}")

            description = " | ".join(description_parts) if description_parts else "Image content"

            logger.info(f"Generated image description: {description[:100]}...")

            return {
                'description': description,
                'labels': labels,
                'objects': objects,
                'ocr_text': ocr_text,
                'has_text': bool(ocr_text),
            }

        except Exception as e:
            logger.error(f"Error describing image with Vision API: {e}")
            return {
                'description': "Image (description unavailable)",
                'labels': [],
                'objects': [],
                'ocr_text': "",
                'has_text': False,
            }

    def upload_image_to_gcs(
        self,
        image: Image.Image,
        document_id: str,
        image_num: int
    ) -> str:
        """Upload image to GCS and return public URL.

        Args:
            image: PIL Image object
            document_id: Document ID
            image_num: Image number in document

        Returns:
            Public GCS URL for the image
        """
        try:
            # Generate unique filename
            img_hash = hashlib.md5(image.tobytes()).hexdigest()[:8]
            blob_name = f"document_images/{document_id}/image_{image_num}_{img_hash}.png"

            bucket = self.storage_client.bucket(self.gcs_bucket)
            blob = bucket.blob(blob_name)

            # Convert to bytes and upload
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)

            blob.upload_from_file(img_byte_arr, content_type='image/png')
            blob.make_public()

            logger.info(f"Uploaded image to GCS: {blob_name}")

            return blob.public_url

        except Exception as e:
            logger.error(f"Error uploading image to GCS: {e}")
            return None

    def parse_multimodal_pdf(
        self,
        pdf_bytes: bytes,
        document_id: str
    ) -> Dict[str, Any]:
        """Parse PDF with text and images.

        Args:
            pdf_bytes: PDF file content as bytes
            document_id: Document ID

        Returns:
            Dictionary with parsed content and images
        """
        # Extract text by page
        pdf_file = io.BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)

        text_by_page = {}
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                text_by_page[page_num] = text

        # Extract and process images
        images = self.extract_images_from_pdf(pdf_bytes)
        processed_images = []

        for image_num, (page_num, image) in enumerate(images):
            # Generate description using Vision API
            image_info = self.describe_image(image)

            # Upload to GCS
            image_url = self.upload_image_to_gcs(image, document_id, image_num)

            processed_images.append({
                'page': page_num + 1,  # 1-indexed
                'image_num': image_num,
                'description': image_info['description'],
                'labels': image_info['labels'],
                'ocr_text': image_info['ocr_text'],
                'has_text': image_info['has_text'],
                'url': image_url,
                'gcs_path': f"document_images/{document_id}/image_{image_num}.png",
                'size': image.size
            })

        # Create page data with image context
        pages_data = []
        for page_num, text in text_by_page.items():
            # Find images on this page
            page_images = [img for img in processed_images if img['page'] == page_num + 1]

            page_data = {
                'page_number': page_num + 1,
                'text': text,
                'images': page_images,
                'has_images': len(page_images) > 0,
                'image_count': len(page_images)
            }

            pages_data.append(page_data)

        logger.info(f"Parsed multimodal PDF: {len(pages_data)} pages, {len(processed_images)} images")

        return {
            'pages': pages_data,
            'total_pages': len(reader.pages),
            'total_images': len(processed_images),
            'images': processed_images,
            'has_images': len(processed_images) > 0
        }

    def parse_multimodal_docx(
        self,
        docx_bytes: bytes,
        document_id: str
    ) -> Dict[str, Any]:
        """Parse DOCX with text and images.

        Args:
            docx_bytes: DOCX file content as bytes
            document_id: Document ID

        Returns:
            Dictionary with parsed content and images
        """
        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)

        # Extract text
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        full_text = "\n\n".join(text_parts)

        # Extract and process images
        images = self.extract_images_from_docx(docx_bytes)
        processed_images = []

        for image_num, (para_idx, image) in enumerate(images):
            # Generate description using Vision API
            image_info = self.describe_image(image)

            # Upload to GCS
            image_url = self.upload_image_to_gcs(image, document_id, image_num)

            processed_images.append({
                'paragraph_index': para_idx,
                'image_num': image_num,
                'description': image_info['description'],
                'labels': image_info['labels'],
                'ocr_text': image_info['ocr_text'],
                'has_text': image_info['has_text'],
                'url': image_url,
                'size': image.size
            })

        logger.info(f"Parsed multimodal DOCX: {len(text_parts)} paragraphs, {len(processed_images)} images")

        return {
            'text': full_text,
            'images': processed_images,
            'total_images': len(processed_images),
            'has_images': len(processed_images) > 0,
            'paragraph_count': len(text_parts)
        }
