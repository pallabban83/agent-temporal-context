"""Document parsing utilities for various file formats."""

import io
from typing import Dict, Any, List
from pypdf import PdfReader
from docx import Document
import pdfplumber
from logging_config import get_logger

logger = get_logger(__name__)


class DocumentParser:
    """Parse various document formats and extract text with table support."""

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file.

        Args:
            file_bytes: PDF file content as bytes

        Returns:
            Extracted text content
        """
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")

            return full_text
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _is_empty_row(row: List[str]) -> bool:
        """Check if a table row contains only empty cells.

        Args:
            row: List of cell values

        Returns:
            True if all cells are empty or None
        """
        return all(not cell or not str(cell).strip() for cell in row)

    @staticmethod
    def _table_to_markdown(table: List[List[str]]) -> str:
        """Convert table data to markdown format with validation and alignment.

        Args:
            table: 2D list of table cells

        Returns:
            Markdown formatted table string, or empty string if invalid
        """
        if not table or len(table) == 0:
            return ""

        # Clean and prepare rows, skip completely empty rows
        rows = []
        for row in table:
            # Skip empty rows
            if DocumentParser._is_empty_row(row):
                continue

            # Replace None with empty string and clean cells
            cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
            rows.append(cleaned_row)

        # Validate table structure
        if len(rows) < 2:  # Need at least header + 1 data row
            logger.warning("Table has fewer than 2 rows, skipping")
            return ""

        # Get max column width for each column
        col_count = max(len(row) for row in rows)

        # Validate minimum columns
        if col_count < 2:  # Need at least 2 columns for meaningful table
            logger.warning("Table has fewer than 2 columns, skipping")
            return ""

        # Pad rows to have consistent column count
        normalized_rows = []
        for row in rows:
            if len(row) < col_count:
                # Pad with empty strings
                padded_row = row + [""] * (col_count - len(row))
                normalized_rows.append(padded_row)
            else:
                normalized_rows.append(row[:col_count])

        # Calculate column widths
        col_widths = [0] * col_count
        for row in normalized_rows:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(cell))

        # Build markdown table
        markdown_lines = []

        # Header row (first row)
        header_cells = [normalized_rows[0][i].ljust(col_widths[i]) for i in range(col_count)]
        header = " | ".join(header_cells)
        markdown_lines.append(f"| {header} |")

        # Separator line
        separator_cells = ["-" * col_widths[i] for i in range(col_count)]
        separator = " | ".join(separator_cells)
        markdown_lines.append(f"| {separator} |")

        # Data rows
        for row in normalized_rows[1:]:
            data_cells = [row[i].ljust(col_widths[i]) for i in range(col_count)]
            data = " | ".join(data_cells)
            markdown_lines.append(f"| {data} |")

        return "\n".join(markdown_lines)

    @staticmethod
    def _extract_text_in_band(page, band_top: float, band_bottom: float, table_bboxes: List[tuple]) -> str:
        """Extract text from a horizontal band on the page, excluding table regions.

        Args:
            page: pdfplumber page object
            band_top: Top y-coordinate of the band
            band_bottom: Bottom y-coordinate of the band
            table_bboxes: List of table bounding boxes to exclude (x0, y0, x1, y1)

        Returns:
            Text content in the band, excluding tables
        """
        # Filter characters that are:
        # 1. Within the vertical band (band_top <= y <= band_bottom)
        # 2. NOT inside any table bbox
        def char_in_band_not_in_tables(char):
            """Check if character is in band and outside all tables."""
            char_x = (char['x0'] + char['x1']) / 2
            char_y = (char['top'] + char['bottom']) / 2

            # Check if in vertical band
            if not (band_top <= char_y <= band_bottom):
                return False

            # Check if NOT in any table
            for bbox in table_bboxes:
                x0, y0, x1, y1 = bbox
                if x0 <= char_x <= x1 and y0 <= char_y <= y1:
                    return False  # Inside table
            return True  # In band and outside all tables

        # Filter the page
        filtered_page = page.filter(char_in_band_not_in_tables)

        # Extract text from filtered page
        text = filtered_page.extract_text() or ""
        return text

    @staticmethod
    def _extract_text_outside_tables(page, table_bboxes: List[tuple]) -> str:
        """Extract text from page excluding table regions.

        Args:
            page: pdfplumber page object
            table_bboxes: List of table bounding boxes (x0, y0, x1, y1)

        Returns:
            Text content outside of table regions
        """
        if not table_bboxes:
            # No tables, extract all text
            return page.extract_text() or ""

        # Filter text chars that are NOT inside any table bbox
        def char_not_in_tables(char):
            """Check if a character is outside all table bounding boxes."""
            char_x = (char['x0'] + char['x1']) / 2
            char_y = (char['top'] + char['bottom']) / 2

            for bbox in table_bboxes:
                x0, y0, x1, y1 = bbox
                if x0 <= char_x <= x1 and y0 <= char_y <= y1:
                    return False  # Inside table
            return True  # Outside all tables

        # Filter the page to exclude table regions
        filtered_page = page.filter(char_not_in_tables)

        # Extract text from filtered page
        text = filtered_page.extract_text() or ""
        return text

    @staticmethod
    def parse_pdf_by_pages(file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from PDF file page by page with table detection.

        This method:
        1. Extracts tables and converts to markdown
        2. Extracts text EXCLUDING table regions (prevents duplication)
        3. Uses global table numbering across all pages

        Args:
            file_bytes: PDF file content as bytes

        Returns:
            Dictionary with page texts, tables, and metadata
        """
        try:
            pdf_file = io.BytesIO(file_bytes)

            page_texts = []
            total_chars = 0
            total_tables = 0
            pages_with_tables = []
            global_table_num = 0  # Global counter for all tables

            # Use pdfplumber for better table extraction
            with pdfplumber.open(pdf_file) as pdf:
                total_pdf_pages = len(pdf.pages)  # Store actual PDF page count

                for page_num, page in enumerate(pdf.pages):
                    # Find tables (returns table objects with both data AND bboxes)
                    table_finder = page.find_tables()

                    # Validate tables and build list with positions
                    validated_tables = []  # List of (y_position, table_markdown, bbox)

                    if table_finder:
                        logger.info(f"Found {len(table_finder)} table(s) on page {page_num + 1}")

                        for table_obj in table_finder:
                            # Validate bbox before processing
                            bbox = table_obj.bbox
                            if not bbox or len(bbox) < 4:
                                logger.warning(
                                    "Skipping table with invalid bbox",
                                    extra={'page': page_num + 1, 'bbox': bbox}
                                )
                                continue

                            # Extract table data from the table object
                            table_data = table_obj.extract()

                            # Convert table to markdown
                            table_markdown = DocumentParser._table_to_markdown(table_data)

                            if table_markdown:
                                # Table passed validation
                                total_tables += 1
                                global_table_num += 1

                                # Add table marker with GLOBAL numbering (spacing handled by join)
                                table_section = f"[TABLE {global_table_num}]\n{table_markdown}\n[END TABLE]"

                                # Store with y-position for sorting
                                y_position = bbox[1]  # y0 (top of table)
                                validated_tables.append((y_position, table_section, bbox))
                            else:
                                # Table failed validation - don't use it
                                logger.info(f"Skipping invalid table on page {page_num + 1} (failed validation)")

                    # Update pages_with_tables only AFTER validation
                    if validated_tables:
                        pages_with_tables.append(page_num + 1)

                    # Extract table bboxes for filtering
                    table_bboxes = [bbox for _, _, bbox in validated_tables]

                    # Sort tables by vertical position (top to bottom)
                    validated_tables.sort(key=lambda x: x[0])

                    # Extract text in segments between tables
                    text_segments = []

                    if not validated_tables:
                        # No tables - extract all text
                        all_text = page.extract_text() or ""
                        if all_text.strip():
                            text_segments.append((0, all_text.strip()))
                    else:
                        # Extract text in bands between tables
                        for i, (y_pos, table_md, bbox) in enumerate(validated_tables):
                            # Define band boundaries
                            if i == 0:
                                # Text before first table
                                band_top = 0
                            else:
                                # Text between previous table and this table
                                band_top = validated_tables[i-1][2][3]  # y1 of previous table

                            band_bottom = bbox[1]  # y0 of current table

                            # Only extract if band is valid (non-negative height)
                            if band_bottom > band_top:
                                # Extract text in this band
                                band_text = DocumentParser._extract_text_in_band(page, band_top, band_bottom, table_bboxes)
                                if band_text.strip():
                                    text_segments.append((band_top, band_text.strip()))
                            else:
                                # Overlapping tables - skip this band
                                logger.debug(f"Skipping overlapping table band on page {page_num + 1} (top={band_top}, bottom={band_bottom})")

                        # Extract text after last table
                        last_table_bottom = validated_tables[-1][2][3]  # y1 of last table
                        after_text = DocumentParser._extract_text_in_band(page, last_table_bottom, page.height, table_bboxes)
                        if after_text.strip():
                            text_segments.append((last_table_bottom, after_text.strip()))

                    # Combine text segments and tables in correct order
                    # Merge lists and sort by y-position
                    all_content = []
                    all_content.extend([(y, 'text', text) for y, text in text_segments])
                    all_content.extend([(y, 'table', md) for y, md, _ in validated_tables])
                    all_content.sort(key=lambda x: x[0])

                    # Build page content maintaining document order
                    page_content_parts = [content for _, _, content in all_content]

                    # Combine text and tables for this page
                    if page_content_parts:
                        page_full_text = "\n\n".join(page_content_parts)
                        page_texts.append(page_full_text)
                        total_chars += len(page_full_text)
                    else:
                        # Empty page - skip entirely (don't add empty string)
                        logger.info(f"Skipping empty page {page_num + 1}")

            logger.info(f"Extracted {total_chars} characters from PDF ({total_pdf_pages} pages, {total_tables} tables)")
            logger.info(f"Non-empty pages: {len(page_texts)}, Pages with tables: {len(pages_with_tables)}")
            logger.info(f"Successfully maintained document order and avoided text duplication")

            return {
                'page_texts': page_texts,
                'total_pages': total_pdf_pages,  # Actual PDF page count
                'non_empty_pages': len(page_texts),  # Pages with content
                'total_chars': total_chars,
                'total_tables': total_tables,
                'pages_with_tables': pages_with_tables,
                'has_tables': total_tables > 0
            }
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file.

        Args:
            file_bytes: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")

            return full_text
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_text(file_bytes: bytes, encoding: str = 'utf-8') -> str:
        """Extract text from plain text file.

        Args:
            file_bytes: Text file content as bytes
            encoding: Text encoding (default: utf-8)

        Returns:
            Decoded text content
        """
        try:
            text = file_bytes.decode(encoding)
            logger.info(f"Decoded {len(text)} characters from text file")
            return text
        except UnicodeDecodeError:
            # Try alternative encodings
            for alt_encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_bytes.decode(alt_encoding)
                    logger.info(f"Decoded {len(text)} characters using {alt_encoding}")
                    return text
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file with common encodings")

    @classmethod
    def parse_document(cls, file_bytes: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """Parse document based on file type.

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Dictionary with extracted text and metadata
        """
        filename_lower = filename.lower()

        # Determine file type and parse accordingly
        if filename_lower.endswith('.pdf') or content_type == 'application/pdf':
            text = cls.parse_pdf(file_bytes)
            doc_type = 'pdf'
        elif filename_lower.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text = cls.parse_docx(file_bytes)
            doc_type = 'docx'
        elif filename_lower.endswith('.doc'):
            raise ValueError("Legacy .doc format not supported. Please convert to .docx or PDF")
        elif filename_lower.endswith(('.txt', '.md', '.markdown')):
            text = cls.parse_text(file_bytes)
            doc_type = 'text'
        else:
            # Try as text file
            try:
                text = cls.parse_text(file_bytes)
                doc_type = 'text'
            except ValueError:
                raise ValueError(f"Unsupported file format: {filename}. Supported formats: PDF, DOCX, TXT, MD")

        return {
            'text': text,
            'type': doc_type,
            'char_count': len(text),
            'word_count': len(text.split())
        }
